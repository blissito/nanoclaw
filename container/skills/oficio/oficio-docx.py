#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
oficio-docx.py — Genera un oficio / carta institucional DESCTI en .docx editable.

Patrón template+data (igual que structured_doc para tabulares): el agente solo
provee un JSON con el CONTENIDO; este helper arma membrete, márgenes, tipografía,
cuerpo justificado y firma. El .docx reflujea solo en Word — NUNCA se corta en la
página 1 (que es el defecto de la rama HTML de página fija) — y queda editable para
el cliente (consecutivo, firma, sello).

Uso:
  python3 oficio-docx.py datos.json /workspace/group/attachments/oficio-x.docx
  cat datos.json | python3 oficio-docx.py - /workspace/group/attachments/oficio-x.docx

Imprime la ruta del .docx generado. Entrégala con
  send_message({ document_path: "<ruta>", text: "..." }).

Requiere: python-docx (baked en la imagen del container tras el rebuild).

JSON esperado (todo opcional salvo `cuerpo`; los campos que falten se omiten):
{
  "no_oficio": "DESCTI/DI/045/2026",        # si falta, deja "_______/2026" para llenar
  "lugar_fecha": "Pachuca de Soto, Hidalgo, a 5 de junio de 2026",
  "destinatario": {
    "nombre": "MTRO. JUAN PÉREZ LÓPEZ",
    "cargo": "Director General de Vinculación y Transferencia",
    "dependencia": "Secretaría de Educación Pública de Hidalgo",
    "linea_extra": "P R E S E N T E."        # default: "P R E S E N T E."
  },
  "asunto": "Solicitud de colaboración interinstitucional",
  "cuerpo": ["Por medio del presente...", "En seguimiento a..."],   # OBLIGATORIO
  "despedida": "Sin otro particular, le reitero la seguridad de mi atenta consideración.",
  "firma": { "nombre": "DR. ...", "cargo": "Director de Innovación", "lema": "..." },
  "ccp": ["C.c.p. Archivo.", "C.c.p. Minutario."],
  "membrete_lineas": ["DIRECCIÓN DE INNOVACIÓN"],   # texto bajo el logo (opcional)
  "logo_path": "/abs/logo.png",             # opcional; si falta baja el logo DESCTI
  "logo_url": "https://..."                 # opcional; override de la URL del logo
}
"""

import json
import os
import sys
import tempfile
import urllib.request

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: falta python-docx. La imagen del container aún no se reconstruyó con la "
        "dependencia.\nMientras tanto genera el oficio con structured_doc (PDF que pagina, "
        "no se corta) y avisa que el .docx queda pendiente del rebuild.\n"
    )
    sys.exit(3)

GRANATE = RGBColor(0x7A, 0x1F, 0x3D)   # paleta institucional DESCTI/Hidalgo
GRIS = RGBColor(0x55, 0x55, 0x55)
FONT = "Arial"
DEFAULT_LOGO_URL = (
    "https://easybits-public.fly.storage.tigris.dev/699f35cbc8ad86037eda62b1/5QW"
)


def _set_font(run, size=11, bold=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # fijar la fuente también en ascii/hAnsi/cs para que no caiga al default de Word
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)


def _bottom_border(paragraph, color="7A1F3D", size="12"):
    """Regla horizontal bajo un párrafo (size en 1/8 pt → 12 = 1.5pt)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _download_logo(url):
    try:
        fd, path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        req = urllib.request.Request(url, headers={"User-Agent": "oficio-docx"})
        with urllib.request.urlopen(req, timeout=20) as resp, open(path, "wb") as out:
            out.write(resp.read())
        return path
    except Exception as exc:  # noqa: BLE001
        sys.stderr.write(
            f"AVISO: no se pudo bajar el logo ({exc}); el oficio saldrá sin membrete gráfico.\n"
        )
        return None


def build(data, out_path):
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # --- membrete en el header (se repite en CADA página) ---
    logo_path = data.get("logo_path")
    if not logo_path:
        logo_path = _download_logo(data.get("logo_url", DEFAULT_LOGO_URL))
    header = sec.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path and os.path.exists(logo_path):
        hp.add_run().add_picture(logo_path, height=Cm(1.8))
    for line in data.get("membrete_lineas", []):
        mp = header.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(mp.add_run(line), size=9, bold=True, color=GRANATE)
    _bottom_border(header.add_paragraph())  # regla granate bajo el membrete

    def para(align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10, space_before=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.line_spacing = 1.15
        return p

    # --- No. de oficio (arriba a la derecha) ---
    no_of = data.get("no_oficio")
    p = para(align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    _set_font(p.add_run("Oficio No. " + (no_of or "_______/2026")), size=11, bold=True)

    # --- lugar y fecha ---
    if data.get("lugar_fecha"):
        p = para(align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=16)
        _set_font(p.add_run(data["lugar_fecha"]), size=11)

    # --- destinatario ---
    dest = data.get("destinatario") or {}
    if dest:
        for key, bold in (("nombre", True), ("cargo", False), ("dependencia", False)):
            if dest.get(key):
                p = para(align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
                _set_font(p.add_run(dest[key]), size=11, bold=bold)
        extra = dest.get("linea_extra", "P R E S E N T E.")
        if extra:
            p = para(align=WD_ALIGN_PARAGRAPH.LEFT, space_after=16, space_before=4)
            _set_font(p.add_run(extra), size=11, bold=True)

    # --- asunto ---
    if data.get("asunto"):
        p = para(align=WD_ALIGN_PARAGRAPH.LEFT, space_after=16)
        _set_font(p.add_run("Asunto: "), size=11, bold=True)
        _set_font(p.add_run(data["asunto"]), size=11)

    # --- cuerpo ---
    cuerpo = data.get("cuerpo") or []
    if isinstance(cuerpo, str):
        cuerpo = [cuerpo]
    for parrafo in cuerpo:
        p = para(space_after=10)
        _set_font(p.add_run(parrafo), size=11)

    # --- despedida ---
    if data.get("despedida"):
        p = para(space_after=28, space_before=8)
        _set_font(p.add_run(data["despedida"]), size=11)

    # --- firma ---
    firma = data.get("firma") or {}
    if firma:
        p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, space_before=24)
        _set_font(p.add_run("A T E N T A M E N T E"), size=11, bold=True)
        if firma.get("lema"):
            p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
            _set_font(p.add_run(firma["lema"]), size=9, color=GRANATE)
        else:
            para(space_after=18)  # espacio en blanco para la rúbrica
        if firma.get("nombre"):
            p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
            _set_font(p.add_run(firma["nombre"]), size=11, bold=True)
        if firma.get("cargo"):
            p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
            _set_font(p.add_run(firma["cargo"]), size=10)

    # --- c.c.p. ---
    ccp = data.get("ccp") or []
    if ccp:
        para(space_after=4, space_before=24)
        for line in ccp:
            p = para(align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
            _set_font(p.add_run(line), size=8, color=GRIS)

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 3:
        sys.stderr.write("Uso: oficio-docx.py <datos.json|-> <salida.docx>\n")
        sys.exit(2)
    src, out = sys.argv[1], sys.argv[2]
    raw = sys.stdin.read() if src == "-" else open(src, encoding="utf-8").read()
    data = json.loads(raw)
    if not data.get("cuerpo"):
        sys.stderr.write("ERROR: el JSON necesita al menos `cuerpo` (lista de párrafos).\n")
        sys.exit(2)
    print(build(data, out))


if __name__ == "__main__":
    main()
